from flask import Flask, render_template, request, redirect, url_for, send_file
import csv, os
from datetime import datetime
import xlsxwriter
from docx import Document

app = Flask(__name__)
DATA_FILE = "data.csv"
UPLOAD_FOLDER = os.path.join("static", "uploads")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def load_data():
    if not os.path.exists(DATA_FILE):
        return []
    with open(DATA_FILE, newline="", encoding="utf-8") as f:
        return list(csv.reader(f))

def save_data(rows):
    with open(DATA_FILE, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerows(rows)

@app.route("/", methods=["GET","POST"])
def index():
    rows = load_data()
    jumlah_hari_ini = sum(1 for r in rows if r[6][:10] == datetime.now().strftime("%Y-%m-%d"))
    sedang_di_dalam = sum(1 for r in rows if r[8] == "Aktif")

    if request.method == "POST":
        nama = request.form["nama"]
        nik = request.form["nik"]
        tujuan = request.form["tujuan"]
        user = request.form["user"]
        asal = request.form["asal"]
        foto = request.files["foto"]

        filename = f"{nik}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.jpg"
        foto.save(os.path.join(UPLOAD_FOLDER, filename))

        masuk = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        rows.append([nama, nik, tujuan, user, asal, filename, masuk, "nan", "Aktif"])
        save_data(rows)
        return redirect(url_for("laporan"))

    return render_template("index.html", jumlah_hari_ini=jumlah_hari_ini, sedang_di_dalam=sedang_di_dalam)

@app.route("/laporan", methods=["GET","POST"])
def laporan():
    rows = load_data()
    keyword = ""
    tanggal = ""

    if request.method == "POST":
        keyword = request.form.get("keyword", "").lower()
        tanggal = request.form.get("tanggal", "")

        # filter keyword (nama/NIK)
        if keyword:
            rows = [r for r in rows if keyword in r[0].lower() or keyword in r[1].lower()]

        # filter tanggal (cek kolom masuk)
        if tanggal:
            rows = [r for r in rows if r[6].startswith(tanggal)]

    return render_template("laporan.html", data=rows, keyword=keyword, tanggal=tanggal)

@app.route("/checkout/<int:index>", methods=["POST"])
def checkout(index):
    rows = load_data()
    if 0 <= index < len(rows):
        rows[index][7] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        rows[index][8] = "Sudah Keluar"
        save_data(rows)
    return redirect(url_for("laporan"))

@app.route("/export_excel")
def export_excel():
    rows = load_data()
    filename = "laporan_tamu.xlsx"
    workbook = xlsxwriter.Workbook(filename)
    worksheet = workbook.add_worksheet()

    headers = ["Nama","NIK","Tujuan","User","Asal","Masuk","Keluar","Status"]
    for col, h in enumerate(headers):
        worksheet.write(0, col, h)

    for row_idx, row in enumerate(rows, start=1):
        data = row[:5] + row[6:]
        for col_idx, val in enumerate(data):
            worksheet.write(row_idx, col_idx, val)

    workbook.close()
    return send_file(filename, as_attachment=True)

@app.route("/export_word")
def export_word():
    rows = load_data()
    doc = Document()
    doc.add_heading("Laporan Tamu", level=1)

    headers = ["Nama","NIK","Tujuan","User","Asal","Masuk","Keluar","Status"]
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h

    for i, row in enumerate(rows, start=1):
        row_cells = table.rows[i].cells
        data = row[:5] + row[6:]
        for j, val in enumerate(data):
            row_cells[j].text = val

    filename = "laporan_tamu.docx"
    doc.save(filename)
    return send_file(filename, as_attachment=True)

@app.route("/download_foto/<filename>")
def download_foto(filename):
    img_path = os.path.join(UPLOAD_FOLDER, filename)
    if os.path.exists(img_path):
        return send_file(img_path, as_attachment=True)
    return "Foto tidak ditemukan", 404

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8000)
